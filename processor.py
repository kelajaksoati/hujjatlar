import os
import openpyxl
from openpyxl.styles import Font
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document

def smart_rename(filename):
    base_name, extension = os.path.splitext(filename)
    # Nomdagi belgilarni tozalash
    clean_name = base_name.replace("_", " ").replace("-", " ").strip()
    
    # Rus maktablari uchun nomni formatlash
    if any(x in clean_name.lower() for x in ["rus", "klass", "рус", "класс"]):
        clean_name = clean_name.replace("rus", "").replace("рус", "").strip()
        final_name = f"RUS_{clean_name.upper()}"
    # BSB/CHSB uchun nomni formatlash
    elif any(x in clean_name.lower() for x in ["bsb", "chsb", "сор", "соч"]):
        name_only = clean_name.lower().replace('bsb', '').replace('chsb', '').strip()
        final_name = f"BSB_CHSB_{name_only.upper()}"
    else:
        final_name = clean_name.title().replace(" ", "_")
        
    return f"@ISH_REJA_UZ_{final_name}{extension.lower()}"

# --- KATEGORIYANI ANIQLASH (Yangi qo'shildi) ---
def get_category_by_name(new_name):
    name_lower = new_name.lower()
    
    # 1. BSB / CHSB / СОР / СОЧ
    if any(x in name_lower for x in ["bsb", "chsb", "сор", "соч"]):
        return "BSB_CHSB"
    
    # 2. Rus maktablari
    if any(x in name_lower for x in ["rus", "klass", "рус", "класс"]):
        return "Rus_maktab"
    
    # 3. Yuqori sinflar
    if any(x in name_lower for x in ["5-","6-","7-","8-","9-","10-","11-"]):
        return "Yuqori"
    
    # 4. Default: Boshlang'ich
    return "Boshlang'ich"

def edit_excel(path):
    try:
        wb = openpyxl.load_workbook(path)
        for ws in wb.worksheets:
            ws.insert_rows(1)
            ws['A1'] = "@ish_reja_uz kanali uchun maxsus tayyorlandi"
            ws['A1'].font = Font(bold=True, color="FF0000", size=11)
        wb.save(path)
    except Exception as e:
        print(f"Excel error: {e}")

def add_pdf_watermark(path):
    try:
        reader = PdfReader(path)
        writer = PdfWriter()
        packet = BytesIO()
        can = canvas.Canvas(packet)
        can.setFont("Helvetica-Bold", 40)
        can.setFillGray(0.5, 0.2)
        can.saveState()
        can.translate(300, 450)
        can.rotate(45)
        can.drawCentredString(0, 0, "@ish_reja_uz")
        can.restoreState()
        can.save()
        packet.seek(0)
        watermark_page = PdfReader(packet).pages[0]
        for page in reader.pages:
            page.merge_page(watermark_page)
            writer.add_page(page)
        with open(path, "wb") as f:
            writer.write(f)
    except Exception as e:
        print(f"PDF error: {e}")

def edit_docx(path):
    try:
        doc = Document(path)
        text = "@ish_reja_uz kanali uchun maxsus tayyorlandi"
        if doc.paragraphs:
            doc.paragraphs[0].insert_paragraph_before(text)
        else:
            doc.add_paragraph(text)
        doc.save(path)
    except Exception as e:
        print(f"Docx error: {e}")
